from django.shortcuts import render, redirect, get_object_or_404
from .models import UserProfile
from .forms import UserProfileForm, WorkExperienceForm, EducationForm, HobbyForm
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from docx import Document
from django.shortcuts import render

def form_view(request, template_id):
    # Burada template_id ile hangi şablon seçildiği bilgisi alınacak.
    # Form gönderildiyse bilgileri kaydet, değilse formu göster.

    if request.method == "POST":
        # Form verilerini kaydet
        # Örneğin: isim, soyisim, eğitim, vs.
        # Kaydetme işlemi burada yapılacak (model tanımlaman lazım)
        pass

    context = {
        "template_id": template_id
    }
    return render(request, "form.html", context)

def home(request):
    return render(request, 'cv_app/home.html')

def create_profile(request):
    if request.method == 'POST':
        form = UserProfileForm(request.POST, request.FILES)
        if form.is_valid():
            profile = form.save()
            return redirect('edit_profile', profile_id=profile.id)
    else:
        form = UserProfileForm()
    return render(request, 'cvapp/create_profile.html', {'form': form})

def edit_profile(request, profile_id):
    profile = get_object_or_404(UserProfile, id=profile_id)

    work_form = WorkExperienceForm(prefix='work')
    education_form = EducationForm(prefix='edu')
    hobby_form = HobbyForm(prefix='hobby')

    if request.method == 'POST':
        if 'add_work' in request.POST:
            work_form = WorkExperienceForm(request.POST, prefix='work')
            if work_form.is_valid():
                work_exp = work_form.save(commit=False)
                work_exp.user_profile = profile
                work_exp.save()
                return redirect('edit_profile', profile_id=profile.id)

        elif 'add_education' in request.POST:
            education_form = EducationForm(request.POST, prefix='edu')
            if education_form.is_valid():
                edu = education_form.save(commit=False)
                edu.user_profile = profile
                edu.save()
                return redirect('edit_profile', profile_id=profile.id)

        elif 'add_hobby' in request.POST:
            hobby_form = HobbyForm(request.POST, prefix='hobby')
            if hobby_form.is_valid():
                hobby = hobby_form.save(commit=False)
                hobby.user_profile = profile
                hobby.save()
                return redirect('edit_profile', profile_id=profile.id)

    context = {
        'profile': profile,
        'work_experiences': profile.work_experiences.all(),
        'educations': profile.educations.all(),
        'hobbies': profile.hobbies.all(),
        'work_form': work_form,
        'education_form': education_form,
        'hobby_form': hobby_form,
    }
    return render(request, 'cvapp/edit_profile.html', context)

def download_pdf(request, profile_id):
    profile = get_object_or_404(UserProfile, id=profile_id)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{profile.name}_cv.pdf"'

    p = canvas.Canvas(response)
    p.drawString(100, 800, f"CV of {profile.name}")

    y = 780
    for work in profile.work_experiences.all():
        p.drawString(100, y, f"Work: {work.position} at {work.company}")
        y -= 20

    for edu in profile.educations.all():
        p.drawString(100, y, f"Education: {edu.degree} at {edu.institution}")
        y -= 20

    for hobby in profile.hobbies.all():
        p.drawString(100, y, f"Hobby: {hobby.name}")
        y -= 20

    p.showPage()
    p.save()
    return response

def download_word(request, profile_id):
    profile = get_object_or_404(UserProfile, id=profile_id)
    document = Document()
    document.add_heading(f'CV of {profile.name}', 0)

    document.add_heading('Work Experiences', level=1)
    for work in profile.work_experiences.all():
        document.add_paragraph(f"{work.position} at {work.company} ({work.start_date} - {work.end_date})")

    document.add_heading('Education', level=1)
    for edu in profile.educations.all():
        document.add_paragraph(f"{edu.degree} at {edu.institution} ({edu.start_date} - {edu.end_date})")

    document.add_heading('Hobbies', level=1)
    for hobby in profile.hobbies.all():
        document.add_paragraph(hobby.name)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{profile.name}_cv.docx"'
    document.save(response)
    return response
