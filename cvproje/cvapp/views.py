from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from .forms import CVForm, EducationForm, ExperienceForm, SkillForm
from .models import CV
from django.contrib import messages
import json

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.template.loader import render_to_string, get_template
from xhtml2pdf import pisa
import io

def index(request):
    return render(request, 'cvapp/index.html', {'templates': CV.TEMPLATE_CHOICES})

def create_cv(request):
    if request.method == 'POST':
        form = CVForm(request.POST, request.FILES)
        if form.is_valid():
            cv = form.save(commit=False)
            
            # Fotoğraf kontrolü ve kaydetme
            if 'photo' in request.FILES:
                cv.photo = request.FILES['photo']
            
            # Eğitim bilgilerini JSON formatında kaydet
            educations = []
            for i in range(int(request.POST.get('education_count', 0))):
                education = {
                    'school': request.POST.get(f'education_{i}_school', ''),
                    'degree': request.POST.get(f'education_{i}_degree', ''),
                    'year': request.POST.get(f'education_{i}_year', '')
                }
                educations.append(education)
            cv.education = json.dumps(educations, ensure_ascii=False)
            
            # Deneyim bilgilerini JSON formatında kaydet
            experiences = []
            for i in range(int(request.POST.get('experience_count', 0))):
                experience = {
                    'company': request.POST.get(f'experience_{i}_company', ''),
                    'position': request.POST.get(f'experience_{i}_position', ''),
                    'duration': request.POST.get(f'experience_{i}_duration', ''),
                    'description': request.POST.get(f'experience_{i}_description', '')
                }
                experiences.append(experience)
            cv.experience = json.dumps(experiences, ensure_ascii=False)
            
            # Beceri bilgilerini JSON formatında kaydet
            skills = []
            for i in range(int(request.POST.get('skill_count', 0))):
                skill = {
                    'name': request.POST.get(f'skill_{i}_name', ''),
                    'level': request.POST.get(f'skill_{i}_level', '')
                }
                skills.append(skill)
            cv.skills = json.dumps(skills, ensure_ascii=False)
            
            cv.save()
            return redirect('preview_cv', cv_id=cv.id)
    else:
        initial_template = request.GET.get('template', 'template1')
        form = CVForm(initial={'template': initial_template})
    
    return render(request, 'cvapp/create_cv.html', {
        'form': form,
        'education_form': EducationForm(),
        'experience_form': ExperienceForm(),
        'skill_form': SkillForm()
    })

def preview_cv(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id)
    
    # Fotoğraf URL'sini oluştur
    if cv.photo:
        cv.photo_url = request.build_absolute_uri(cv.photo.url)
    
    # Şablon numarasını al (template1 -> 1)
    template_number = cv.template.replace('template', '')
    
    # Şablon rengini belirle
    if cv.template == 'template1':
        cv.template_color = '#0d6efd'
        cv.template_gradient = 'linear-gradient(90deg, #0d6efd, #0a58ca)'
    elif cv.template == 'template2':
        cv.template_color = '#198754'
        cv.template_gradient = 'linear-gradient(90deg, #198754, #146c43)'
    else:
        cv.template_color = '#6f42c1'
        cv.template_gradient = 'linear-gradient(90deg, #6f42c1, #59359a)'
    
    # JSON verilerini düzenle
    try:
        if cv.education and cv.education.strip():
            education_data = json.loads(cv.education)
            cv.education = '\n'.join([f"• {edu['school']} - {edu['degree']} ({edu['year']})" for edu in education_data])
        else:
            cv.education = ""
    except json.JSONDecodeError:
        cv.education = ""
    
    try:
        if cv.experience and cv.experience.strip():
            experience_data = json.loads(cv.experience)
            cv.experience = '\n\n'.join([f"• {exp['company']} - {exp['position']}\n  {exp['duration']}\n  {exp['description']}" for exp in experience_data])
        else:
            cv.experience = ""
    except json.JSONDecodeError:
        cv.experience = ""
    
    try:
        if cv.skills and cv.skills.strip():
            skills_data = json.loads(cv.skills)
            cv.skills = '\n'.join([f"• {skill['name']} ({skill['level']})" for skill in skills_data])
        else:
            cv.skills = ""
    except json.JSONDecodeError:
        cv.skills = ""
    
    return render(request, 'cvapp/preview_cv.html', {'cv': cv})

def cv_list(request):
    cvs = CV.objects.all().order_by('-created_at')
    
    # Her CV için JSON verilerini düzenle ve fotoğraf URL'sini oluştur
    for cv in cvs:
        # Fotoğraf URL'sini oluştur
        if cv.photo and hasattr(cv.photo, 'url'):
            cv.photo_url = request.build_absolute_uri(cv.photo.url)
        else:
            cv.photo_url = None
        
        try:
            if cv.education and cv.education.strip():
                education_data = json.loads(cv.education)
                cv.education = '\n'.join([f"• {edu['school']} - {edu['degree']} ({edu['year']})" for edu in education_data])
            else:
                cv.education = ""
        except json.JSONDecodeError:
            cv.education = ""
        
        try:
            if cv.experience and cv.experience.strip():
                experience_data = json.loads(cv.experience)
                cv.experience = '\n\n'.join([f"• {exp['company']} - {exp['position']}\n  {exp['duration']}\n  {exp['description']}" for exp in experience_data])
            else:
                cv.experience = ""
        except json.JSONDecodeError:
            cv.experience = ""
        
        try:
            if cv.skills and cv.skills.strip():
                skills_data = json.loads(cv.skills)
                cv.skills = '\n'.join([f"• {skill['name']} ({skill['level']})" for skill in skills_data])
            else:
                cv.skills = ""
        except json.JSONDecodeError:
            cv.skills = ""
    
    return render(request, 'cvapp/cv_list.html', {'cvs': cvs})

def format_json_data(json_data):
    try:
        data = json.loads(json_data)
        if isinstance(data, list):
            formatted_text = ""
            for item in data:
                if isinstance(item, dict):
                    formatted_text += "\n".join(f"{value}" for value in item.values() if value) + "\n\n"
            return formatted_text.strip()
        return json_data
    except:
        return json_data

def download_cv_pdf(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id)
    
    # Fotoğraf URL'sini oluştur
    if cv.photo:
        cv.photo_url = request.build_absolute_uri(cv.photo.url)
    
    # JSON verilerini formatla
    education = format_json_data(cv.education)
    experience = format_json_data(cv.experience)
    skills = format_json_data(cv.skills)
    
    # Şablon numarasını al (template1 -> 1)
    template_number = cv.template.replace('template', '')
    
    template_name = f'template{template_number}_pdf.html'
    template = get_template(f'cvapp/templates/{template_name}')
    
    context = {
        'cv': cv,
        'education': education,
        'experience': experience,
        'skills': skills
    }
    
    html = template.render(context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{cv.name}_CV.pdf"'
    
    # PDF oluşturma seçenekleri
    pdf_options = {
        'encoding': 'UTF-8',
        'page-size': 'A4',
        'margin-top': '20mm',
        'margin-right': '20mm',
        'margin-bottom': '20mm',
        'margin-left': '20mm',
        'enable-local-file-access': True
    }
    
    pisa_status = pisa.CreatePDF(html, dest=response, **pdf_options)
    if pisa_status.err:
        return HttpResponse('PDF oluşturulurken bir hata oluştu')
    return response

def download_cv_word(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id)
    
    # Fotoğraf URL'sini oluştur
    if cv.photo:
        cv.photo_url = request.build_absolute_uri(cv.photo.url)
    
    # JSON verilerini formatla
    education = format_json_data(cv.education)
    experience = format_json_data(cv.experience)
    skills = format_json_data(cv.skills)
    
    # Şablon numarasını al (template1 -> 1)
    template_number = cv.template.replace('template', '')
    
    template_name = f'template{template_number}_word.html'
    template = get_template(f'cvapp/templates/{template_name}')
    
    context = {
        'cv': cv,
        'education': education,
        'experience': experience,
        'skills': skills
    }
    
    html = template.render(context)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{cv.name}_CV.docx"'
    
    # HTML'i Word belgesine dönüştür
    doc = Document()
    
    # Başlık ve kişisel bilgiler
    heading = doc.add_heading(cv.name, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Şablon rengini belirle
    if cv.template == 'template1':
        color = RGBColor(13, 110, 253)  # #0d6efd
    elif cv.template == 'template2':
        color = RGBColor(25, 135, 84)   # #198754
    else:
        color = RGBColor(111, 66, 193)  # #6f42c1
    
    # Başlık rengini ayarla
    for run in heading.runs:
        run.font.color.rgb = color
    
    # Fotoğraf ekle
    if cv.photo:
        try:
            doc.add_picture(cv.photo.path, width=Inches(1.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    
    # İletişim bilgileri
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f'Email: {cv.email}\n')
    contact.add_run(f'Telefon: {cv.phone}')
    
    # Eğitim
    education_heading = doc.add_heading('Eğitim', level=1)
    education_heading.runs[0].font.color.rgb = color
    doc.add_paragraph(education)
    
    # Deneyim
    experience_heading = doc.add_heading('Deneyim', level=1)
    experience_heading.runs[0].font.color.rgb = color
    doc.add_paragraph(experience)
    
    # Beceriler
    skills_heading = doc.add_heading('Beceriler', level=1)
    skills_heading.runs[0].font.color.rgb = color
    doc.add_paragraph(skills)
    
    # Belgeyi kaydet
    doc.save(response)
    return response

def delete_cv(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id)
    cv_name = cv.name
    cv.delete()
    messages.success(request, f'{cv_name} isimli CV başarıyla silindi.')
    return redirect('cv_list')
